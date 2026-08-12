"""
generate_whitepaper.py -- Build the CUBE v5 Scientific & Technical White Paper (.docx).

Usage:
    "C:\\Users\\param\\anaconda3\\envs\\CUBE\\python.exe" generate_whitepaper.py

Source material: README.md, CUBE_GUIDE.md, GROUP_PREDICTOR_REFERENCE.md (this repo).
"""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

BASE = Path(__file__).parent
FLOWCHART_PNG = Path(r"C:\Users\param\AppData\Local\Temp\claude\d--CUBE\ffe2605d-f1f0-4a98-92c7-5f84f3bc3cbb\scratchpad\pipeline_comparison.png")

# ── Palette ──────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x36, 0x5D)
TEAL   = RGBColor(0x00, 0x80, 0x80)
ICE    = "F4F7F9"
ICE_RGB = RGBColor(0xF4, 0xF7, 0xF9)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x1B, 0x27, 0x33)
GREY   = RGBColor(0x5B, 0x65, 0x70)
TABLE_ALT = "EDF2F5"

# ── low-level helpers ────────────────────────────────────────────────────

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hexcolor)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        node = OxmlElement(f'w:{tag}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)


def para_shading(para, hexcolor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hexcolor)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def para_border(para, sides=("top", "bottom", "left", "right"), color="1B365D", sz="18"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)


def set_col_widths(table, widths_in):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_in):
            row.cells[i].width = Inches(w)
    for i, w in enumerate(widths_in):
        table.columns[i].width = Inches(w)


def add_run(para, text, bold=False, italic=False, size=None, color=None, font=None):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if font:
        r.font.name = font
    return r


# ── document-level styling ───────────────────────────────────────────────

def setup_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    # Header/footer: running title + page numbers
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    r = hp.add_run("CUBE v5 White Paper  |  Valiathan, Karolinska Institutet")
    r.font.size = Pt(7.5)
    r.font.color.rgb = GREY
    r.italic = True
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


def add_heading(doc, text, level=1, space_before=10, space_after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.keep_with_next = True
    sizes = {1: 15, 2: 12.5, 3: 11}
    colors = {1: NAVY, 2: TEAL, 3: NAVY}
    r = add_run(para, text, bold=True, size=sizes.get(level, 11), color=colors.get(level, NAVY))
    if level == 1:
        para_border(para, sides=("bottom",), color="008080", sz="14")
        para.paragraph_format.space_after = Pt(6)
    if level == 2:
        para.paragraph_format.space_before = Pt(8)
    return para


def add_body(doc, text, size=10, space_after=5, justify=True):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    if justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _render_inline(para, text, size=size)
    return para


def _render_inline(para, text, size=10, color=DARK, bold_default=False):
    """Minimal **bold** and *italic* inline markdown support."""
    import re
    tokens = re.split(r'(\*\*.+?\*\*|\*.+?\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            add_run(para, tok[2:-2], bold=True, size=size, color=color)
        elif tok.startswith('*') and tok.endswith('*') and len(tok) > 1:
            add_run(para, tok[1:-1], italic=True, size=size, color=color)
        else:
            add_run(para, tok, bold=bold_default, size=size, color=color)


def add_bullets(doc, items, size=9.5):
    for it in items:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.left_indent = Inches(0.2)
        _render_inline(para, it, size=size)


def add_callout(doc, title, text, accent="008080"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, ICE)
    set_cell_margins(cell, top=90, bottom=90, left=160, right=160)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:start')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '30'); left.set(qn('w:color'), accent)
    borders.append(left)
    for side in ('top', 'end', 'bottom'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tcPr.append(borders)
    cell.paragraphs[0].clear() if cell.paragraphs[0].runs else None
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    add_run(p1, title, bold=True, size=9.5, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    _render_inline(p2, text, size=9.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, header_row, data_rows, widths=None, header_bg="1B365D",
              font_size=8.7, header_size=8.7):
    n_cols = len(header_row)
    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tbl.rows[0].cells
    for i, htext in enumerate(header_row):
        set_cell_bg(hdr_cells[i], header_bg)
        set_cell_margins(hdr_cells[i])
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, htext, bold=True, size=header_size, color=WHITE)
    for ridx, row in enumerate(data_rows):
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i])
            if ridx % 2 == 0:
                set_cell_bg(cells[i], TABLE_ALT)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            _render_inline(p, str(val), size=font_size, color=DARK)
    if widths:
        set_col_widths(tbl, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)


# ── Build document ────────────────────────────────────────────────────────

def build():
    doc = Document()
    setup_doc(doc)

    # ── Title block ──
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(title_p, "CUBE v5 — Comprehensive Unsupervised Behavioral Explorer",
            bold=True, size=20, color=NAVY)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(6)
    add_run(sub_p, "Architecture, Algorithmic Evolution, Multivariate Predictive Modeling, "
                   "and Positioning within the Unsupervised Behavioral-Analysis Literature",
            italic=True, size=12, color=TEAL)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(10)
    add_run(meta_p, "P. Valiathan, MBB", bold=True, size=10, color=DARK)
    add_run(meta_p, "  \u00b7  Karolinska Institutet  \u00b7  Technical White Paper  \u00b7  v5 Release",
            size=10, color=GREY)
    para_border(meta_p, sides=("bottom",), color="008080", sz="10")

    # ── Abstract ──
    add_body(doc,
        "CUBE is a six-step, GUI-integrated pipeline that extends the B-SOiD unsupervised "
        "pose-behavior-classification framework (Hsu & Yttri, 2021, *Nat. Commun.* 12:5188) into a "
        "single, reproducible tool spanning multi-camera video acquisition through group-level "
        "statistical inference. Where the original B-SOiD method leaves feature engineering, cluster "
        "validation, temporal smoothing, video-based annotation, and statistics as separate, "
        "researcher-assembled steps, CUBE unifies them behind one desktop application with automated "
        "quality gating at every stage. This paper documents CUBE's architecture, the specific "
        "algorithmic modifications that raise clustering fidelity above baseline B-SOiD, the "
        "multivariate Group Predictor's generalized-linear-model (GLM) classification framework, and "
        "CUBE's position relative to VAME, keypoint-MoSeq, SimBA, and JAABA.", size=9.7)

    add_page_break(doc)

    # ── 1. Architecture & Workflow Evolution ──
    add_heading(doc, "1. Architectural and Workflow Evolution", level=1)
    add_body(doc,
        "Legacy B-SOiD-based analysis is a chain of independently operated tools: a researcher "
        "manually splits raw video, runs DeepLabCut (DLC) in its own GUI or notebook, hand-formats the "
        "resulting CSV/H5 pose files, feeds them into a standalone B-SOiD script with a single fixed "
        "feature scale and one UMAP/HDBSCAN run, and finally exports cluster labels to external "
        "spreadsheets or statistical software for group comparison. Every hand-off is a place where "
        "formatting drifts, quality is unchecked, and results become difficult to reproduce. CUBE "
        "replaces this chain with six in-process pipeline steps, all driven from one `customtkinter` "
        "desktop application (`cube.py`) that lazy-loads the two standalone GUIs (`cube_analyser.py`, "
        "`cube_video_explorer.py`) via `importlib` only when the user reaches those steps — keeping "
        "startup light while letting either tool also run independently for ad-hoc re-analysis.", size=9.7)

    add_table(doc,
        ["Step", "Stage", "What happens"],
        [
            ["0", "Acquisition sync", "Bonsai 4-camera recording synchronization with per-frame fill-frame detection; outputs corrected, timestamp-aligned MP4s per camera."],
            ["1", "Pose estimation", "2D: batch SuperAnimal-quadruped DLC inference with Smart Adapt. 3D: per-camera DLC Zoo adaptation \u2192 aniposelib RANSAC triangulation \u2192 4-coordinate (x,y,z,likelihood) H5, plus an NVENC-encoded quad-camera composite video for downstream QC and clip review."],
            ["2", "Pre-processing", "Bodypart conservation and confidence-gated filtering (`cube_core.py`); produces a per-session bodypart-confidence report before any features are computed."],
            ["3", "Clustering engine", "Multi-scale feature extraction \u2192 UMAP \u2192 HDBSCAN DBCV sweep \u2192 consensus/co-association clustering \u2192 bidirectional split/merge refinement \u2192 MLP classifier \u2192 Baum-Welch HMM temporal smoothing."],
            ["4", "Annotation", "Interactive Video Explorer: cluster labeling against embedded example-clip playback, run standalone or in-pipeline."],
            ["5", "Analysis", "Behaviour Analyser and Group Predictor: ethograms, group statistics, transition dynamics, reclustering, and multivariate group classification."],
        ],
        widths=[0.4, 1.1, 5.0])

    add_callout(doc, "Reproducibility by construction",
        "Pipeline state autosaves to `.pipeline_session.json` after every completed step (paths, "
        "parameters, completion flags, DLC model state — not the numerical data itself, which is "
        "re-loaded from disk on resume). A crashed or interrupted run reopens to the exact step it "
        "left off, which the legacy manual chain has no equivalent for.")

    add_heading(doc, "Usability and desktop modularity", level=2)
    add_bullets(doc, [
        "**Lazy-loading architecture** — `cube_analyser.py` and `cube_video_explorer.py` are imported via `importlib.util.spec_from_file_location` only when their step is opened, and both also execute as fully standalone scripts against any correctly-structured output folder.",
        "**Automated quality diagnostics** — `cluster_validity.png` renders the HDBSCAN condensed tree with a silhouette diagram so cluster hierarchy is inspectable rather than opaque; `split_merge_refinement.png` shows a before/after contingency heatmap confirming exactly which clusters were split or merged and why.",
        "**Burned-in video audit overlays** — labeled videos carry a live `C{id}` cluster badge and, when triggered, an amber \u201cTURNED AWAY\u201d banner, so detection and classification quality are visually auditable without a separate script.",
    ])

    # ── 2. Multivariate Behavioral Fingerprinting ──
    add_heading(doc, "2. Multivariate Behavioral Fingerprinting and the Group Predictor GLM", level=1)
    add_body(doc,
        "Per-cluster tests (Kruskal-Wallis, Dunn's) answer *which single behaviour differs*, but miss "
        "combinations of individually modest differences that jointly separate experimental groups. "
        "The Group Predictor addresses this by training a multivariate classifier per animal's full "
        "behavioural fingerprint, evaluated across three parallel, independently informative feature "
        "spaces:", size=9.7)

    add_table(doc,
        ["Dimension", "Definition", "What a positive result implies"],
        [
            ["Frequency", "Bouts of a cluster per second of session", "Manipulation changes how *often* a behaviour is initiated"],
            ["Total Duration", "Cumulative time in a cluster as a fraction of session length", "Manipulation changes how *much* time is spent in a behaviour, independent of bout count"],
            ["Transition Probability", "Row-normalized, off-diagonal cluster\u2192cluster switching matrix (self-persistence diagonal excluded to decouple sequencing from duration)", "Manipulation changes the behavioural *grammar* — what follows what — even if frequency/duration are unchanged"],
        ],
        widths=[1.15, 2.9, 2.45])

    add_heading(doc, "Preprocessing pipeline (identical across all three models)", level=2)
    add_body(doc,
        "Raw per-animal feature matrices pass through a fixed five-stage scikit-learn pipeline before "
        "any classifier sees them:", size=9.7, space_after=3)
    add_table(doc,
        ["#", "Stage", "Purpose"],
        [
            ["1", "`VarianceThreshold` (< 10\u207b\u00b9\u2070)", "Drops clusters with effectively identical counts across every animal — pure noise dimensions."],
            ["2", "`RobustScaler`", "Centers/scales on median and IQR rather than mean/SD, so one outlier animal cannot dominate normalization."],
            ["3", "`PolynomialFeatures`", "Frequency/Duration models only — adds pairwise interaction terms (e.g. freq(C4)\u00d7freq(C5)) to capture \u201chigh-C4-AND-high-C5-together\u201d synergy; skipped for the Transition model, which already encodes co-occurrence structurally."],
            ["4", "Adaptive PCA", "Triggered only when n_features > n_animals; compresses to min(n_animals\u22122, 15) components, with feature importances back-projected to the original cluster space so interpretation is not lost to the rotation."],
            ["5", "Classifier", "Elastic-Net multinomial logistic regression (default) or linear SVM with Platt-scaled probabilities."],
        ],
        widths=[0.3, 1.6, 4.6])

    add_heading(doc, "Why a generalized linear model, and why Elastic Net specifically", level=2)
    add_body(doc,
        "The default classifier is a **multinomial Elastic-Net logistic regression** — a GLM with a "
        "logit link and a combined L1/L2 penalty. **Lasso (L1)** drives redundant cluster coefficients "
        "to exactly zero, performing automatic feature selection; **Ridge (L2)** stabilizes that "
        "selection against the correlated-cluster problem endemic to B-SOiD-style output, where L1 "
        "alone would arbitrarily pick one of several co-occurring clusters and zero the rest, making "
        "feature importances unstable run-to-run. For \u22653 groups the model fits a multinomial softmax "
        "rather than one-vs-rest, which is more principled under unequal group sizes. Hyperparameters "
        "(C \u2208 {0.001\u201310}, L1-ratio \u2208 {0.3, 0.5, 0.7}, widened for n\u226525 animals) are tuned by inner "
        "cross-validation with `min(3, min_group_size)` folds, and `class_weight=\"balanced\"` "
        "compensates for uneven cohort sizes. A linear SVM with Platt scaling is offered as an "
        "alternative when linear separability with a hard margin is preferred.", size=9.7)

    add_callout(doc, "Why Random Forest was deliberately excluded",
        "With the cohort sizes typical of behavioural pharmacology (n = 5\u201310 animals per group), "
        "a 100-tree forest fit on 8\u201310 animals is overparameterized by roughly an order of magnitude "
        "— it memorizes individual training animals rather than learning generalizable structure, "
        "producing inflated LOO accuracy. Random Forest becomes an appropriate choice only once "
        "cohorts exceed roughly n > 30 per group; below that, the low-variance, easily-regularized "
        "linear GLM is the more honest estimator.")

    add_heading(doc, "Statistical validation layer", level=2)
    add_bullets(doc, [
        "**Leave-One-Out Cross-Validation (LOO-CV)** — with n < 30 animals, LOO maximizes training data per fold (N\u22121 animals) and is the standard choice at typical CUBE cohort sizes.",
        "**Stratified permutation testing** (Phipson & Smyth, 2010) — group labels are shuffled and the full LOO procedure re-run to build a null distribution; p = (permutations \u2265 observed + 1) / (n_permutations + 1) avoids a p = 0 artifact. `StratifiedShuffleSplit` preserves group proportions in each null fold. Conditional (fast, fixed feature selection) and nested (re-runs feature selection per permutation; gold-standard for publication) variants are both available.",
        "**Balanced accuracy and Cohen's \u03ba** — balanced accuracy (mean per-class recall) is the headline statistic, since raw accuracy is misleading under group-size imbalance; \u03ba > 0.6 combined with a significant permutation p-value is treated as a strong indicator of a genuine, replicable effect.",
        "**Shapley feature importances** — exact enumeration of all 2\u1d3a coalitions for \u2264 8 clusters, Monte Carlo sampling (150 orderings) above that; unlike raw GLM coefficients, Shapley \u03c6 values are fair, order-independent attributions robust to the correlated-feature problem that destabilizes coefficient-only interpretation.",
    ])

    add_page_break(doc)

    # ── 3. Paradigm comparison ──
    add_heading(doc, "3. Pipeline Paradigm Comparison — Classic DLC\u2192B-SOiD vs. CUBE v5", level=1)
    add_body(doc,
        "The two workflows below are contrasted for the common case of a single-camera (2D) recording. "
        "The classic route requires five distinct tools and manual data hand-offs at every arrow; each "
        "hand-off is an unaudited opportunity for format drift, silent quality loss, or lost provenance. "
        "CUBE's six steps run inside one application with an explicit confidence-gated checkpoint "
        "(Step 2) before any clustering occurs — pose data that fails bodypart-conservation or "
        "confidence criteria is flagged before it can silently degrade the embedding.", size=9.7, space_after=6)

    doc.add_picture(str(FLOWCHART_PNG), width=Inches(6.75))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    add_run(cap, "Figure 1. Classic DeepLabCut\u2192B-SOiD workflow (left) vs. the CUBE v5 integrated "
                 "pipeline (right), 2D single-camera example. Grey arrows denote manual, offline hand-offs "
                 "between separate tools; navy/teal arrows denote automated in-process data flow.",
            italic=True, size=8.3, color=GREY)

    # ── 4. Algorithmic improvements over baseline B-SOiD ──
    add_heading(doc, "4. Algorithmic Improvements Over Baseline B-SOiD", level=1)
    add_body(doc,
        "CUBE keeps B-SOiD's core architecture — frame/bin-wise pose features embedded with UMAP and "
        "partitioned with HDBSCAN — because it is fast (minutes, not hours) and needs no GPU training "
        "loop. What it changes is everything around that core that determines whether the resulting "
        "clusters are trustworthy:", size=9.7)

    add_table(doc,
        ["Stage", "Baseline B-SOiD", "CUBE v5 modification", "Why it raises fidelity"],
        [
            ["Feature extraction", "Single fixed temporal scale", "Multi-scale (100 ms + 200 ms, +50 ms above 60 fps) with within-bin positional variance and temporal-lag drift features", "Separates sustained postures from rapid oscillatory motion (tremor, flinching) that a single scale conflates or misses entirely"],
            ["Dimensionality control", "Not applied", "Adaptive PCA pre-reduction, auto-triggered when n_features \u2265 n_samples/5", "Prevents nearest-neighbour graph degradation in high-dimensional feature spaces (common once body-region weighting expands feature count)"],
            ["Turn-away/occlusion handling", "Not addressed — turned-away frames embed as if real postural variation", "Explicit `detect_turned_away_bins` (Head/Mouth region confidence + nose-keypoint likelihood, sustained-window debounced) excludes flagged bins from training and assigns a dedicated \u201cTurned Away\u201d label", "Prevents an artefact of camera occlusion from fragmenting or contaminating genuine behaviour clusters"],
            ["UMAP neighbourhood scaling", "Fixed/manual `n_neighbors`", "fps- and session-length-adaptive `n_neighbors` formula, floor raised 15\u219230 after real-data seed-stability testing (mean pairwise ARI 0.31\u21920.57)", "Small neighbourhoods make UMAP's stochastic optimization highly seed-sensitive; the raised floor materially improves embedding reproducibility"],
            ["HDBSCAN parameterization", "Single manual `min_cluster_size`", "40-step sweep across `min_cluster_size`, both `eom`/`leaf` methods, selection by DBCV (`relative_validity_`)", "Removes a hand-tuned, dataset-specific hyperparameter and replaces it with a data-driven quality criterion"],
            ["Cluster-count selection / consensus", "One HDBSCAN fit, one seed", "Optional consensus/co-association clustering across `consensus_n_seeds` (default 8) independent UMAP+HDBSCAN runs, auto-triggered when seed-sweep mean ARI falls below threshold; Ward-linkage clustering of the resulting co-association matrix", "Resolves seed-instability by construction rather than by guessing why seeds disagree; validated 3\u201314\u00d7 higher within- vs. between-cluster co-association"],
            ["Fragmentation/over-splitting", "Not addressed", "Bidirectional split/merge refinement — locally re-embeds low-silhouette (\u201cimpure\u201d) clusters, merges sibling clusters barely separated in the condensed tree, bounded by candidate/performance caps", "Directly targets the documented B-SOiD failure mode of one behaviour fragmenting across several clusters (e.g. \u201clicking\u201d across 3, \u201csniffing\u201d across 7 in validation data)"],
            ["Temporal smoothing", "None — raw frame-wise labels flicker", "`hmmlearn.GaussianHMM` (default since Aug 2026) fitted on the MLP's per-bin class-probability vectors rather than collapsed hard labels, trained at 100ms-bin resolution, with a per-cluster self-transition prior derived from each cluster's own mean bout duration (`hmmlearn.CategoricalHMM` on hard frame-repeated labels with one flat prior remains available via `hmm_emission_mode=\"categorical\"`), Viterbi decoding", "Eliminates single-frame label flicker without discarding the underlying frame-level classification; down-weights smoothing confidence on frames the classifier itself was unsure about, and avoids over-smoothing naturally brief behaviours under the same transition assumption used for long ones"],
            ["Noise/confidence filtering", "Not addressed at the feature level", "Adaptive visibility/occlusion feature block plus automatic confidence-based bodypart down-weighting (session-fraction-based taper, not mean- or max-based, after both were shown to under- or over-correct on real data)", "Chronically unreliable bodyparts stop contributing flat, near-identical feature vectors that otherwise push HDBSCAN noise up and DBCV toward degeneracy"],
        ],
        widths=[1.05, 1.55, 2.6, 2.3], font_size=8.0, header_size=8.3)

    add_callout(doc, "Net effect on usefulness relative to baseline B-SOiD",
        "None of the above changes the fundamental B-SOiD architecture or its runtime profile — every "
        "addition is off-by-default or auto-triggered only when a diagnostic (seed-sweep ARI, DBCV "
        "degeneracy) indicates it is needed. The result is a tool that keeps B-SOiD's speed advantage "
        "while closing the specific gaps — fragmentation, occlusion artefacts, seed instability, and "
        "absence of temporal structure — that otherwise required moving to a heavier architecture "
        "such as VAME or keypoint-MoSeq to fix.")

    # ── 5. Literature benchmarking ──
    add_heading(doc, "5. Literature Positioning", level=1)
    add_body(doc,
        "CUBE's cluster-quality additions are best understood as targeted, cheap fixes for the same "
        "failure modes that motivate heavier architectures elsewhere in the field, applied inside the "
        "existing fast pipeline rather than through a full re-embedding or generative rewrite.", size=9.7)

    add_table(doc,
        ["Framework", "Core approach", "How CUBE compares"],
        [
            ["B-SOiD\n(Hsu & Yttri 2021)", "Frame-wise pose features \u2192 UMAP \u2192 HDBSCAN; no GPU training loop; minutes per run", "CUBE's base architecture. Resolves B-SOiD's own documented weaknesses (frame fragmentation, occlusion artefacts) via HMM Viterbi smoothing, condensed-tree split/merge passes, and adaptive visibility features, without changing runtime class."],
            ["VAME", "RNN autoencoder learns a temporal embedding before clustering, so sequence structure is captured pre-clustering", "CUBE achieves comparable sequence-level smoothing (condensed-tree merge pass + optional body-region weighting) without VAME's GPU retraining loop or re-embedding architecture — a lighter-weight approximation of the same diagnosis, not a re-implementation."],
            ["keypoint-MoSeq", "Generative, uncertainty-aware model that automatically down-weights unreliable keypoints", "CUBE's adaptive visibility/occlusion features and automatic confidence-based bodypart weighting are a simpler, discriminative analogue: low per-bodypart confidence becomes an explicit feature axis / weight taper rather than a modeled latent uncertainty, letting the existing HDBSCAN step isolate unreliable bins."],
            ["SimBA / JAABA", "Supervised classifiers trained on manually labeled behaviour exemplars", "CUBE (like base B-SOiD) is fully unsupervised at the discovery stage — zero manual labeling is required to obtain the initial cluster set; Step 4's Video Explorer adds *interpretive* labels after the fact, which is optional and does not feed back into re-training the clustering."],
        ],
        widths=[1.1, 2.55, 3.85], font_size=8.3)

    add_page_break(doc)

    # ── 6. Experimental design & statistics ──
    add_heading(doc, "6. Experimental Design and Statistical Framework", level=1)
    add_body(doc,
        "The Analyser supports up to three independent metadata label columns per animal. Label 3 "
        "doubles as an **Animal ID**, which is what enables repeated-measures pairing across "
        "timepoints — this single design choice lets the same infrastructure serve both between-"
        "subjects pharmacology studies and within-subject longitudinal designs without separate code "
        "paths.", size=9.7)

    add_table(doc,
        ["Design", "Omnibus test", "Post-hoc", "Correction", "Notes"],
        [
            ["Independent Groups (default)", "Kruskal-Wallis", "Dunn's test (Mann-Whitney U fallback without `scikit-posthocs`)", "Benjamini-Hochberg FDR, pooled across every cluster/panel reported together", "Two-part decomposition splits **prevalence** (Fisher's exact / \u03c7\u00b2 — is the behaviour present at all?) from **magnitude** (present-only Kruskal-Wallis) via a `sig_driver` label \u2208 {magnitude, prevalence, both, none}"],
            ["Repeated Measures", "Wilcoxon signed-rank (2 levels) or Friedman's test (\u22653 levels)", "Pairwise Wilcoxon signed-rank", "Same FDR pooling as Independent Groups", "Pairing via Animal ID (Label 3); only animals with a matching, non-blank ID at *every* level are included (\u22653 matched minimum); `sig_driver` reports \u201cn/a\u201d since the structural-zero decomposition is Independent-Groups-only"],
        ],
        widths=[1.4, 1.5, 1.5, 1.5, 2.6], font_size=8.0, header_size=8.3)

    add_body(doc,
        "A third \u201cMixed Design\u201d (independent groups each followed across time) is not yet available. "
        "The Group Predictor's LOO-CV and permutation framework (\u00a72) operates as a complementary, "
        "multivariate layer on top of this per-cluster statistical family — the two are intended to be "
        "read together, since a manipulation can shift joint behavioural structure detectably even when "
        "no single cluster survives FDR correction in isolation.", size=9.7)

    # ── 7. Conclusions ──
    add_heading(doc, "7. Conclusions and Deployment Recommendations", level=1)
    add_body(doc,
        "CUBE v5 is most valuable wherever a lab currently runs DLC and B-SOiD as separate tools and "
        "pays the reproducibility cost of manual hand-offs between them. Specific deployment "
        "guidance:", size=9.7)

    add_bullets(doc, [
        "**High-throughput drug/behavioural screening (n = 5\u201330 animals/group):** the Group Predictor's LOO-CV + permutation framework and Elastic-Net GLM are calibrated for exactly this cohort-size regime, where Random Forest and other high-capacity classifiers overfit.",
        "**3D kinematic / multi-camera free-movement ethology:** Step 1's Anipose RANSAC triangulation and Step 3's dimensionality-agnostic feature/UMAP/HDBSCAN/HMM stack (2D and 3D share the same downstream code) make CUBE a direct fit for 4-camera open-field or home-cage rigs without a separate analysis codebase per modality.",
        "**Studies at risk from occlusion or camera turn-away artefacts** (e.g. freely-moving rodents, unrestrained head orientation): the dedicated turned-away detection and exclusion pipeline directly targets a failure mode that silently contaminates clusters in unmodified B-SOiD.",
        "**Longitudinal/repeated-measures pharmacology:** the Animal-ID-based Wilcoxon/Friedman pathway and Group Predictor together give both single-behaviour and joint-fingerprint views of a within-subject effect across timepoints.",
        "**When to stay with a heavier architecture instead:** if seed-sweep ARI remains low even after consensus clustering, or if the behavioural repertoire is dominated by continuous, non-discrete transitions that resist any hard clustering, a full generative re-embedding (keypoint-MoSeq) or temporal-autoencoder approach (VAME) remains the more principled next step — CUBE's mitigations are diagnosed-and-targeted fixes, not a categorical replacement for those architectures.",
    ])

    add_callout(doc, "Bottom line",
        "CUBE v5 converts a five-tool, manually-stitched B-SOiD workflow into one reproducible, "
        "quality-gated application spanning acquisition through group-level statistical inference — "
        "without sacrificing B-SOiD's core runtime advantage over GPU-trained temporal or generative "
        "alternatives.")

    out_path = BASE / "CUBE_v5_White_Paper.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
