"""
md_to_docx.py — Convert CUBE markdown documentation to formatted Word documents.

Converts:
    README.md                    → README.docx
    CUBE_GUIDE.md                → CUBE_GUIDE.docx
    GROUP_PREDICTOR_REFERENCE.md → GROUP_PREDICTOR_REFERENCE.docx

Usage:
    "C:\\Users\\param\\anaconda3\\envs\\CUBE\\python.exe" md_to_docx.py

Requires python-docx (auto-installed on first run):
    pip install python-docx
"""

import re
import sys
from pathlib import Path

# ── Dependency check ──────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    print("python-docx not found — installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


# ── Colour palette ────────────────────────────────────────────────────────────
C_H1           = RGBColor(0x1F, 0x49, 0x7D)  # dark navy
C_H2           = RGBColor(0x2E, 0x74, 0xB5)  # medium blue
C_H3           = RGBColor(0x1F, 0x49, 0x7D)  # dark navy (same as H1)
C_H4           = RGBColor(0x40, 0x40, 0x40)  # dark grey
C_CODE_FG      = RGBColor(0xC0, 0x39, 0x2B)  # deep red  (inline code)
C_LINK         = RGBColor(0x00, 0x56, 0x99)  # link blue
C_TABLE_HDR_BG = RGBColor(0x2E, 0x74, 0xB5)  # header row fill
C_TABLE_HDR_FG = RGBColor(0xFF, 0xFF, 0xFF)  # header row text
C_TABLE_ALT    = RGBColor(0xEE, 0xF3, 0xFA)  # alternate row fill


# ── XML helpers ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, rgb: RGBColor):
    """Set a table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), str(rgb))   # RGBColor.__str__ returns 'RRGGBB'
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def _set_para_border_bottom(para):
    """Add a thin bottom border to a paragraph (used as horizontal rule)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_para_shading(para, hex_fill: str):
    """Set paragraph background shading (used for code blocks)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_fill)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


# ── Inline markdown parser ────────────────────────────────────────────────────

# HTML entity map
_HTML_ENTITIES = {
    '&nbsp;': ' ', '&amp;': '&', '&lt;': '<', '&gt;': '>',
    '&quot;': '"', '&#39;': "'", '&mdash;': '—', '&ndash;': '–',
    '&bull;': '•', '&hellip;': '…',
}

def _decode_entities(text: str) -> str:
    for entity, char in _HTML_ENTITIES.items():
        text = text.replace(entity, char)
    return text


def _strip_inline_html(text: str) -> str:
    """Remove HTML tags, replacing block-level ones with newlines."""
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    return _decode_entities(text)


# Token regex — order matters: bold before italic, to avoid ** matching as two *
_INLINE_RE = re.compile(
    r'(\*\*(.+?)\*\*)'                    # group 1,2  : **bold**
    r'|(\*([^*\n]+?)\*)'                  # group 3,4  : *italic*
    r'|(`([^`\n]+?)`)'                    # group 5,6  : `code`
    r'|(\[([^\]\n]+?)\]\(([^)\n]+?)\))'  # group 7,8,9: [text](url)
    r'|(<b>(.+?)</b>)'                    # group 10,11: <b>bold</b>
    r'|(<i>(.+?)</i>)'                    # group 12,13: <i>italic</i>
    r'|(<a\s[^>]*>(.+?)</a>)',            # group 14,15: <a href>text</a>
    re.DOTALL,
)


def _add_inline_runs(para, text: str, default_bold=False, default_italic=False,
                     default_size_pt: float = None):
    """Parse inline markdown/HTML in *text* and add formatted runs to *para*."""
    text = _decode_entities(text)
    pos = 0
    for m in _INLINE_RE.finditer(text):
        # Plain text before this match
        if pos < m.start():
            run = para.add_run(text[pos:m.start()])
            run.bold = default_bold
            run.italic = default_italic
            if default_size_pt:
                run.font.size = Pt(default_size_pt)

        if m.group(1):      # **bold**
            run = para.add_run(m.group(2))
            run.bold = True
            run.italic = default_italic
        elif m.group(3):    # *italic*
            run = para.add_run(m.group(4))
            run.bold = default_bold
            run.italic = True
        elif m.group(5):    # `code`
            run = para.add_run(m.group(6))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = C_CODE_FG
        elif m.group(7):    # [text](url)
            run = para.add_run(m.group(8))
            run.font.color.rgb = C_LINK
            run.underline = True
        elif m.group(10):   # <b>bold</b>
            run = para.add_run(m.group(11))
            run.bold = True
        elif m.group(12):   # <i>italic</i>
            run = para.add_run(m.group(13))
            run.italic = True
        elif m.group(14):   # <a href>text</a>
            run = para.add_run(m.group(15))
            run.font.color.rgb = C_LINK
            run.underline = True

        if default_size_pt:
            for run in para.runs[-1:]:
                if not run.font.size:
                    run.font.size = Pt(default_size_pt)
        pos = m.end()

    # Remaining plain text
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.bold = default_bold
        run.italic = default_italic
        if default_size_pt:
            run.font.size = Pt(default_size_pt)


# ── Block renderers ───────────────────────────────────────────────────────────

def _render_heading(doc, level: int, text: str):
    """Add a heading paragraph at the given level with CUBE colour scheme."""
    # Strip markdown links and inline code from heading text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = _strip_inline_html(text).strip()

    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
    color_map = {1: C_H1, 2: C_H2, 3: C_H3, 4: C_H4}
    size_map  = {1: 16, 2: 14, 3: 12, 4: 11}

    para = doc.add_paragraph(style=style_map.get(level, 'Heading 4'))
    para.paragraph_format.space_before = Pt({1: 14, 2: 10, 3: 8, 4: 6}.get(level, 6))
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.font.color.rgb = color_map.get(level, C_H4)
    run.font.size = Pt(size_map.get(level, 11))
    run.bold = True
    return para


def _render_code_block(doc, lines: list):
    """Add a fenced code block with monospace font and grey background."""
    if not lines:
        return
    code_text = '\n'.join(lines)
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.25)
    para.paragraph_format.right_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(6)
    _set_para_shading(para, 'F0F0F0')
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)


def _render_table(doc, raw_rows: list):
    """Parse markdown table rows and render as a Word table."""
    parsed = []
    for line in raw_rows:
        line = line.strip()
        # Skip separator rows like |---|:---:|---|
        if re.match(r'^[\|\s\-:]+$', line):
            continue
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')]
        # Remove the empty strings created by leading/trailing pipes
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if cells:
            parsed.append(cells)

    if not parsed:
        return

    n_cols = max(len(r) for r in parsed)
    parsed = [r + [''] * (n_cols - len(r)) for r in parsed]  # pad short rows

    tbl = doc.add_table(rows=len(parsed), cols=n_cols)
    tbl.style = 'Table Grid'

    for row_i, row_data in enumerate(parsed):
        is_header = (row_i == 0)
        word_row  = tbl.rows[row_i]
        for col_i, cell_text in enumerate(row_data):
            cell = word_row.cells[col_i]
            cell.paragraphs[0].clear()
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            _add_inline_runs(para, cell_text, default_bold=is_header,
                             default_size_pt=9.5)
            if is_header:
                for run in para.runs:
                    run.font.color.rgb = C_TABLE_HDR_FG
                    run.bold = True
                _set_cell_bg(cell, C_TABLE_HDR_BG)
            elif row_i % 2 == 0:
                _set_cell_bg(cell, C_TABLE_ALT)

    doc.add_paragraph()  # breathing room after table


def _render_bullet(doc, text: str, indent_level: int = 0):
    """Add a bullet list item."""
    style = 'List Bullet 2' if indent_level > 0 else 'List Bullet'
    try:
        para = doc.add_paragraph(style=style)
    except Exception:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent_level)
    _add_inline_runs(para, text, default_size_pt=10.5)


def _render_numbered(doc, text: str, indent_level: int = 0):
    """Add a numbered list item."""
    style = 'List Number 2' if indent_level > 0 else 'List Number'
    try:
        para = doc.add_paragraph(style=style)
    except Exception:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent_level)
    _add_inline_runs(para, text, default_size_pt=10.5)


def _render_blockquote(doc, text: str):
    """Add a blockquote paragraph."""
    try:
        para = doc.add_paragraph(style='Intense Quote')
    except Exception:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.4)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    _add_inline_runs(para, text, default_italic=True)


def _render_hr(doc):
    """Add a horizontal rule (bottom-bordered blank paragraph)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    _set_para_border_bottom(para)


def _render_paragraph(doc, text: str, centered: bool = False):
    """Add a normal body paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    if centered:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_inline_runs(para, text)


# ── HTML block handling ───────────────────────────────────────────────────────

def _handle_html_line(doc, line: str):
    """
    Extract useful content from single-line HTML tags.
    Multi-line HTML blocks are handled by the caller merging lines first.
    """
    text = line.strip()

    # <img .../>  — skip (can't embed image without the actual file at hand)
    if re.match(r'<img\b', text, re.IGNORECASE):
        return

    # <h1 align="center">...</h1>
    m = re.match(r'<h(\d)[^>]*>(.+?)</h\d>', text, re.IGNORECASE)
    if m:
        _render_heading(doc, int(m.group(1)), _strip_inline_html(m.group(2)).strip())
        return

    # <p align="center">...</p>  or any centered paragraph
    m = re.match(r'<p\b[^>]*align=["\']center["\'][^>]*>(.*?)</p>', text, re.IGNORECASE | re.DOTALL)
    if m:
        inner = _strip_inline_html(m.group(1)).strip()
        if inner:
            _render_paragraph(doc, inner, centered=True)
        return

    # Generic <p>...</p>
    m = re.match(r'<p[^>]*>(.*?)</p>', text, re.IGNORECASE | re.DOTALL)
    if m:
        inner = _strip_inline_html(m.group(1)).strip()
        if inner:
            _render_paragraph(doc, inner)
        return

    # Bare closing/opening tags with no content — skip
    if re.match(r'^</?\w+[^>]*>$', text):
        return

    # Anything else with tags — strip and render as plain text
    plain = _strip_inline_html(text)
    if plain.strip():
        _render_paragraph(doc, plain.strip())


# ── Main converter ────────────────────────────────────────────────────────────

def convert_md_to_docx(md_path: Path, docx_path: Path):
    """Convert a single markdown file to a formatted Word document."""
    print(f"  Converting {md_path.name} ...", end='', flush=True)

    text  = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin   = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Default body font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ── State machine ─────────────────────────────────────────────────────
    i            = 0
    in_code      = False
    code_lines   = []
    in_table     = False
    table_lines  = []
    in_html      = False
    html_lines   = []

    def flush_table():
        nonlocal in_table, table_lines
        if table_lines:
            _render_table(doc, table_lines)
        in_table, table_lines = False, []

    def flush_html():
        nonlocal in_html, html_lines
        combined = ' '.join(html_lines).strip()
        if combined:
            _handle_html_line(doc, combined)
        in_html, html_lines = False, []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Inside fenced code block ─────────────────────────────────────
        if in_code:
            if stripped.startswith('```'):
                _render_code_block(doc, code_lines)
                code_lines = []
                in_code    = False
            else:
                code_lines.append(line)
            i += 1
            continue

        # ── Opening fence ────────────────────────────────────────────────
        if stripped.startswith('```'):
            flush_table()
            flush_html()
            in_code = True
            i += 1
            continue

        # ── HTML multi-line block detection ───────────────────────────────
        # Accumulate lines that contain unmatched HTML tags
        is_html_line = bool(re.match(r'\s*<[a-zA-Z/!]', line))
        if is_html_line:
            flush_table()
            html_lines.append(stripped)
            # Check if the opening tag is closed on this line
            # Simple heuristic: if tag is self-closed or has matching close tag, flush
            tag_opens  = len(re.findall(r'<(?!/)(?!!)[a-zA-Z]', stripped))
            tag_closes = len(re.findall(r'</[a-zA-Z]', stripped)) + len(re.findall(r'/>', stripped))
            if tag_opens <= tag_closes:
                flush_html()
            else:
                in_html = True
            i += 1
            continue
        elif in_html:
            html_lines.append(stripped)
            tag_opens  = sum(len(re.findall(r'<(?!/)(?!!)[a-zA-Z]', l)) for l in html_lines)
            tag_closes = sum(len(re.findall(r'</[a-zA-Z]', l)) + len(re.findall(r'/>', l)) for l in html_lines)
            if tag_opens <= tag_closes:
                flush_html()
            i += 1
            continue

        # ── Horizontal rule (must check before heading underline) ─────────
        if re.match(r'^---+\s*$', stripped):
            flush_table()
            _render_hr(doc)
            i += 1
            continue

        # ── Heading ───────────────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            flush_table()
            _render_heading(doc, len(m.group(1)), m.group(2).strip())
            i += 1
            continue

        # ── Table row ────────────────────────────────────────────────────
        if '|' in stripped and stripped.startswith('|'):
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            flush_table()

        # ── Blank line ────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────
        if stripped.startswith('> '):
            quote_text = stripped[2:]
            # Collect continuation lines starting with >
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('> '):
                i += 1
                quote_text += ' ' + lines[i].strip()[2:]
            _render_blockquote(doc, quote_text.strip())
            i += 1
            continue

        # ── Bullet list ───────────────────────────────────────────────────
        m = re.match(r'^(\s*)([-*+])\s+(.*)', line)
        if m:
            indent_level = len(m.group(1)) // 2
            _render_bullet(doc, m.group(3), indent_level)
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────
        m = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if m:
            indent_level = len(m.group(1)) // 2
            _render_numbered(doc, m.group(2), indent_level)
            i += 1
            continue

        # ── Regular paragraph (with soft-wrap joining) ────────────────────
        para_lines = [stripped]
        while i + 1 < len(lines):
            nxt = lines[i + 1]
            nxt_stripped = nxt.strip()
            # Stop collecting if next line starts a new block element
            if (not nxt_stripped
                    or nxt_stripped.startswith('#')
                    or nxt_stripped.startswith('```')
                    or (nxt_stripped.startswith('|') and '|' in nxt_stripped)
                    or re.match(r'^\s*[-*+]\s', nxt)
                    or re.match(r'^\s*\d+\.\s', nxt)
                    or nxt_stripped.startswith('> ')
                    or re.match(r'^---+\s*$', nxt_stripped)
                    or re.match(r'\s*<[a-zA-Z/!]', nxt)):
                break
            i += 1
            para_lines.append(nxt_stripped)

        para_text = ' '.join(para_lines)
        if para_text:
            _render_paragraph(doc, para_text)

        i += 1

    # Flush any unclosed blocks
    flush_table()
    flush_html()
    if code_lines:
        _render_code_block(doc, code_lines)

    doc.save(str(docx_path))
    size_kb = docx_path.stat().st_size // 1024
    print(f" done  ({size_kb} KB -> {docx_path.name})")


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    base = Path(__file__).parent

    targets = [
        ("README.md",                    "README.docx"),
        ("CUBE_GUIDE.md",                "CUBE_GUIDE.docx"),
        ("GROUP_PREDICTOR_REFERENCE.md", "GROUP_PREDICTOR_REFERENCE.docx"),
    ]

    print("CUBE -- Markdown to Word converter")
    print("=" * 40)
    any_converted = False
    for md_name, docx_name in targets:
        md_path   = base / md_name
        docx_path = base / docx_name
        if md_path.exists():
            try:
                convert_md_to_docx(md_path, docx_path)
                any_converted = True
            except Exception as exc:
                print(f"  ERROR converting {md_name}: {exc}")
        else:
            print(f"  SKIP  {md_name} not found")

    if any_converted:
        print("\nAll done. Open the .docx files in Word to review formatting.")
    else:
        print("\nNo markdown files found to convert.")


if __name__ == "__main__":
    main()
